#!/usr/bin/env python3
"""
Create NDIS and Aged Care Document Templates
Generates professional DOCX and XLSX templates for download
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Warning: openpyxl not installed. Excel templates will be skipped.")

def add_table_border(table):
    """Add borders to a table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_ndis_service_agreement():
    """Create NDIS Service Agreement Template"""
    doc = Document()

    # Title
    title = doc.add_heading('NDIS SERVICE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run('This Service Agreement sets out the terms and conditions under which ').font.size = Pt(11)
    intro.add_run('[Provider Name]').bold = True
    intro.add_run(' will provide NDIS supports to ').font.size = Pt(11)
    intro.add_run('[Participant Name]').bold = True
    intro.add_run('.').font.size = Pt(11)

    doc.add_paragraph()

    # Section 1: Provider Details
    doc.add_heading('1. PROVIDER DETAILS', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)
    table.autofit = False
    table.allow_autofit = False

    provider_fields = [
        ('Provider Name:', '[Enter Provider Name]'),
        ('Trading Name:', '[If different from above]'),
        ('ABN:', '[Enter ABN]'),
        ('NDIS Registration Number:', '[Enter Registration Number]'),
        ('Contact Person:', '[Enter Contact Name]'),
        ('Email/Phone:', '[Enter Email and Phone]')
    ]

    for i, (label, value) in enumerate(provider_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 2: Participant Details
    doc.add_heading('2. PARTICIPANT DETAILS', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)

    participant_fields = [
        ('Participant Name:', '[Enter Participant Name]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('NDIS Number:', '[Enter NDIS Number]'),
        ('Plan Management Type:', '[NDIA Managed / Plan Managed / Self Managed]'),
        ('Plan Manager (if applicable):', '[Enter Plan Manager Details]'),
        ('Emergency Contact:', '[Name and Phone Number]')
    ]

    for i, (label, value) in enumerate(participant_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 3: Services to be Provided
    doc.add_heading('3. SERVICES TO BE PROVIDED', 1)
    table = doc.add_table(rows=5, cols=4)
    add_table_border(table)

    headers = ['Support Type', 'Support Item Number', 'Frequency', 'Rate']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 5):
        table.rows[i].cells[0].text = '[e.g., Assistance with Daily Life]'
        table.rows[i].cells[1].text = '[e.g., 01_011_0107_1_1]'
        table.rows[i].cells[2].text = '[e.g., 10 hours/week]'
        table.rows[i].cells[3].text = '[Rate per hour]'

    doc.add_paragraph()

    # Section 4: Service Delivery
    doc.add_heading('4. SERVICE DELIVERY', 1)
    delivery_items = [
        'Services will be delivered at: [Location(s)]',
        'Service delivery hours: [Days and times]',
        'Services will commence on: [Start Date]',
        'This agreement is valid until: [End Date] or until the participant\'s plan is reviewed'
    ]
    for item in delivery_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Section 5: Participant Rights
    doc.add_heading('5. PARTICIPANT RIGHTS AND RESPONSIBILITIES', 1)
    doc.add_paragraph('As an NDIS participant, you have the right to:')
    rights = [
        'Be treated with dignity and respect',
        'Make choices and decisions about your supports',
        'Privacy and confidentiality of your information',
        'Raise concerns or complaints without fear of retribution',
        'Cancel or change services with reasonable notice'
    ]
    for right in rights:
        p = doc.add_paragraph(right, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Section 6: Provider Responsibilities
    doc.add_heading('6. PROVIDER RESPONSIBILITIES', 1)
    responsibilities = [
        'Deliver supports that align with your NDIS plan and goals',
        'Ensure all staff are appropriately qualified and screened',
        'Maintain appropriate insurance coverage',
        'Report incidents in accordance with NDIS Commission requirements',
        'Provide invoices within agreed timeframes',
        'Protect your privacy and handle information in accordance with privacy laws'
    ]
    for resp in responsibilities:
        p = doc.add_paragraph(resp, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Section 7: Fees and Payment
    doc.add_heading('7. FEES AND PAYMENT', 1)
    payment_text = doc.add_paragraph()
    payment_text.add_run('All fees are in accordance with the current NDIS Pricing Arrangements and Price Limits. Payment terms are ').font.size = Pt(11)
    payment_text.add_run('[e.g., 14 days]').bold = True
    payment_text.add_run(' from the date of invoice.').font.size = Pt(11)

    doc.add_paragraph('Cancellation Policy: [e.g., 48 hours notice required or cancellation fee may apply]')

    doc.add_paragraph()

    # Section 8: Complaints and Feedback
    doc.add_heading('8. COMPLAINTS AND FEEDBACK', 1)
    doc.add_paragraph('If you have any concerns or complaints, please contact:')
    doc.add_paragraph('[Complaints Officer Name and Contact Details]')
    doc.add_paragraph('You can also contact the NDIS Commission on 1800 035 544 or visit www.ndiscommission.gov.au')

    doc.add_paragraph()

    # Section 9: Termination
    doc.add_heading('9. TERMINATION OF AGREEMENT', 1)
    doc.add_paragraph('Either party may terminate this agreement by providing [e.g., 14 days] written notice. The provider must continue to deliver supports during the notice period unless otherwise agreed.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    doc.add_heading('10. SIGNATURES', 1)

    sig_table = doc.add_table(rows=6, cols=2)
    add_table_border(sig_table)

    sig_table.rows[0].cells[0].text = 'PARTICIPANT (or Representative)'
    sig_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(sig_table.rows[0].cells[0], 'D9D9D9')
    sig_table.rows[0].cells[1].text = 'PROVIDER'
    sig_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    set_cell_background(sig_table.rows[0].cells[1], 'D9D9D9')

    sig_table.rows[1].cells[0].text = 'Name: [Print Name]'
    sig_table.rows[1].cells[1].text = 'Name: [Print Name]'

    sig_table.rows[2].cells[0].text = 'Signature:\n\n'
    sig_table.rows[2].cells[1].text = 'Signature:\n\n'

    sig_table.rows[3].cells[0].text = 'Date: [DD/MM/YYYY]'
    sig_table.rows[3].cells[1].text = 'Date: [DD/MM/YYYY]'

    sig_table.rows[4].cells[0].text = 'Relationship to Participant:\n[If signing on behalf]'
    sig_table.rows[4].cells[1].text = 'Position: [Position Title]'

    sig_table.rows[5].cells[0].text = 'Authority to sign:\n[If applicable]'
    sig_table.rows[5].cells[1].text = 'Company: [Provider Name]'

    return doc

def create_ndis_incident_report():
    """Create NDIS Incident Report Form"""
    doc = Document()

    # Title
    title = doc.add_heading('NDIS INCIDENT REPORT FORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Important Notice
    notice = doc.add_paragraph()
    notice.add_run('IMPORTANT: ').bold = True
    notice.add_run('This form must be completed for any incident involving an NDIS participant, including injuries, abuse, neglect, medication errors, property damage, or service delivery issues. Report to the NDIS Commission within required timeframes.')
    notice.paragraph_format.space_after = Pt(12)

    # Section 1: Incident Details
    doc.add_heading('1. INCIDENT DETAILS', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    incident_fields = [
        ('Date of Incident:', '[DD/MM/YYYY]'),
        ('Time of Incident:', '[HH:MM AM/PM]'),
        ('Location of Incident:', '[Specific location]'),
        ('Type of Incident:', '[ ] Injury  [ ] Abuse/Neglect  [ ] Medication Error\n[ ] Property Damage  [ ] Other: _______'),
        ('Severity:', '[ ] Minor  [ ] Moderate  [ ] Severe  [ ] Critical'),
        ('Was emergency service contacted?', '[ ] Yes  [ ] No   If yes, which service: _______'),
        ('Incident Reference Number:', '[Auto-generated or assigned]')
    ]

    for i, (label, value) in enumerate(incident_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 2: Participant Information
    doc.add_heading('2. PARTICIPANT INFORMATION', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    participant_fields = [
        ('Participant Name:', '[Full Name]'),
        ('NDIS Number:', '[NDIS Number]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('Contact Details:', '[Phone/Email]')
    ]

    for i, (label, value) in enumerate(participant_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 3: Reporter Information
    doc.add_heading('3. PERSON REPORTING INCIDENT', 1)
    table = doc.add_table(rows=5, cols=2)
    add_table_border(table)

    reporter_fields = [
        ('Reporter Name:', '[Full Name]'),
        ('Position/Role:', '[Position]'),
        ('Contact Number:', '[Phone]'),
        ('Email:', '[Email]'),
        ('Date Reported:', '[DD/MM/YYYY]')
    ]

    for i, (label, value) in enumerate(reporter_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 4: Incident Description
    doc.add_heading('4. DETAILED DESCRIPTION OF INCIDENT', 1)
    doc.add_paragraph('Provide a detailed factual description of what happened (who, what, when, where, why, how):')

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    cell = table.rows[0].cells[0]
    # Add multiple line breaks for space
    cell.text = '\n\n\n\n\n\n\n\n'

    doc.add_paragraph()

    # Section 5: Witnesses
    doc.add_heading('5. WITNESSES (If applicable)', 1)
    table = doc.add_table(rows=3, cols=3)
    add_table_border(table)

    headers = ['Witness Name', 'Contact Details', 'Relationship to Participant']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 3):
        table.rows[i].cells[0].text = ''
        table.rows[i].cells[1].text = ''
        table.rows[i].cells[2].text = ''

    doc.add_paragraph()

    # Section 6: Immediate Actions
    doc.add_heading('6. IMMEDIATE ACTIONS TAKEN', 1)
    doc.add_paragraph('Describe actions taken immediately following the incident:')

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    cell = table.rows[0].cells[0]
    cell.text = '\n\n\n\n\n'

    doc.add_paragraph()

    # Section 7: Injuries or Impact
    doc.add_heading('7. INJURIES OR IMPACT ON PARTICIPANT', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    impact_fields = [
        ('Was the participant injured?', '[ ] Yes  [ ] No'),
        ('Type of injury (if applicable):', '[Describe injury]'),
        ('Medical treatment required?', '[ ] Yes  [ ] No   If yes, describe: _______'),
        ('Emotional/psychological impact:', '[Describe any distress or impact]')
    ]

    for i, (label, value) in enumerate(impact_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 8: Notifications
    doc.add_heading('8. NOTIFICATIONS MADE', 1)

    notifications = [
        'Participant/Family notified: [ ] Yes  [ ] No   Date/Time: _______',
        'NDIS Commission notified: [ ] Yes  [ ] No   Date/Time: _______',
        'Plan Manager notified: [ ] Yes  [ ] No   Date/Time: _______',
        'Police notified: [ ] Yes  [ ] No   Date/Time: _______',
        'Other (specify): [ ] Yes  [ ] No   Who: _______  Date/Time: _______'
    ]

    for notification in notifications:
        doc.add_paragraph(notification, style='List Bullet')

    doc.add_paragraph()

    # Section 9: Follow-up Actions
    doc.add_heading('9. FOLLOW-UP ACTIONS REQUIRED', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    cell = table.rows[0].cells[0]
    cell.text = '\n\n\n\n\n'

    doc.add_paragraph()

    # Section 10: Sign-off
    doc.add_heading('10. MANAGER REVIEW AND SIGN-OFF', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    signoff_fields = [
        ('Reviewed by (Name):', '[Manager Name]'),
        ('Position:', '[Position Title]'),
        ('Signature:', '\n\n'),
        ('Date:', '[DD/MM/YYYY]')
    ]

    for i, (label, value) in enumerate(signoff_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Footer note
    footer_note = doc.add_paragraph()
    footer_note.add_run('Note: ').bold = True
    footer_note.add_run('Keep this form confidential and store securely for 7 years. Reportable incidents must be submitted to the NDIS Commission within 24 hours (for serious incidents) or 5 business days (for other incidents).')

    return doc

def create_ndis_support_plan():
    """Create NDIS Support Plan Template"""
    doc = Document()

    # Title
    title = doc.add_heading('MY NDIS SUPPORT PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Participant Info
    doc.add_heading('PARTICIPANT INFORMATION', 1)
    table = doc.add_table(rows=5, cols=2)
    add_table_border(table)

    fields = [
        ('Name:', '[Participant Name]'),
        ('NDIS Number:', '[NDIS Number]'),
        ('Plan Start Date:', '[DD/MM/YYYY]'),
        ('Plan End Date:', '[DD/MM/YYYY]'),
        ('Plan Management Type:', '[ ] NDIA Managed  [ ] Plan Managed  [ ] Self Managed')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # My Goals
    doc.add_heading('MY GOALS', 1)
    doc.add_paragraph('These are the goals I want to achieve during this plan period:')

    for i in range(1, 4):
        doc.add_heading(f'Goal {i}:', 2)
        table = doc.add_table(rows=4, cols=2)
        add_table_border(table)

        goal_fields = [
            ('My Goal:', '[Describe your goal]'),
            ('Why this goal matters to me:', '[Explain why this is important]'),
            ('How I will achieve this:', '[Steps or supports needed]'),
            ('How I will know I\'ve achieved it:', '[Success indicators]')
        ]

        for j, (label, value) in enumerate(goal_fields):
            row = table.rows[j]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            set_cell_background(row.cells[0], 'E7E6E6')
            row.cells[1].text = value

        doc.add_paragraph()

    # Core Supports Budget
    doc.add_heading('CORE SUPPORTS BUDGET', 1)
    table = doc.add_table(rows=5, cols=3)
    add_table_border(table)

    headers = ['Support Category', 'Budget Amount', 'How I will use it']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    core_categories = [
        'Assistance with Daily Life',
        'Transport',
        'Consumables',
        'Assistance with Social & Community Participation'
    ]

    for i, category in enumerate(core_categories, 1):
        table.rows[i].cells[0].text = category
        table.rows[i].cells[1].text = '$[Amount]'
        table.rows[i].cells[2].text = '[Description]'

    doc.add_paragraph()

    # Capacity Building Budget
    doc.add_heading('CAPACITY BUILDING BUDGET', 1)
    table = doc.add_table(rows=4, cols=3)
    add_table_border(table)

    headers = ['Support Category', 'Budget Amount', 'How I will use it']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    capacity_categories = [
        'Support Coordination',
        'Improved Living Arrangements',
        'Increased Social & Community Participation'
    ]

    for i, category in enumerate(capacity_categories, 1):
        table.rows[i].cells[0].text = category
        table.rows[i].cells[1].text = '$[Amount]'
        table.rows[i].cells[2].text = '[Description]'

    doc.add_paragraph()

    # Current Providers
    doc.add_heading('MY CURRENT PROVIDERS', 1)
    table = doc.add_table(rows=4, cols=4)
    add_table_border(table)

    headers = ['Provider Name', 'Service Type', 'Contact Details', 'Frequency']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 4):
        table.rows[i].cells[0].text = '[Provider Name]'
        table.rows[i].cells[1].text = '[Service]'
        table.rows[i].cells[2].text = '[Contact]'
        table.rows[i].cells[3].text = '[Frequency]'

    doc.add_paragraph()

    # Plan Review
    doc.add_heading('PLAN REVIEW', 1)
    doc.add_paragraph('My plan will be reviewed on: [DD/MM/YYYY]')
    doc.add_paragraph('To prepare for my review, I will track:')

    tracking = [
        'Progress towards my goals',
        'What supports are working well',
        'What supports are not working',
        'Any changes in my circumstances or needs',
        'New goals I want to work towards'
    ]

    for item in tracking:
        doc.add_paragraph(item, style='List Bullet')

    return doc

def create_aged_care_plan():
    """Create Aged Care Individual Care Plan Template"""
    doc = Document()

    # Title
    title = doc.add_heading('INDIVIDUAL CARE PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Client Details
    doc.add_heading('CLIENT DETAILS', 1)
    table = doc.add_table(rows=8, cols=2)
    add_table_border(table)

    fields = [
        ('Client Name:', '[Full Name]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('Address:', '[Residential Address]'),
        ('Contact Number:', '[Phone]'),
        ('Emergency Contact:', '[Name and Phone]'),
        ('Care Plan Start Date:', '[DD/MM/YYYY]'),
        ('Care Plan Review Date:', '[DD/MM/YYYY]'),
        ('GP Details:', '[Name, Clinic, Phone]')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Health Status
    doc.add_heading('HEALTH STATUS & MEDICAL CONDITIONS', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    cell = table.rows[0].cells[0]
    cell.text = 'List all current medical conditions, diagnoses, and relevant health history:\n\n\n\n\n'

    doc.add_paragraph()

    # Current Medications
    doc.add_heading('CURRENT MEDICATIONS', 1)
    table = doc.add_table(rows=5, cols=5)
    add_table_border(table)

    headers = ['Medication Name', 'Dosage', 'Frequency', 'Route', 'Prescriber']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 5):
        for j in range(5):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Functional Assessment
    doc.add_heading('FUNCTIONAL ASSESSMENT', 1)

    table = doc.add_table(rows=11, cols=3)
    add_table_border(table)

    headers = ['Activity', 'Level of Assistance', 'Notes/Equipment Used']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    activities = [
        'Personal Hygiene',
        'Bathing/Showering',
        'Dressing',
        'Mobility',
        'Transfers',
        'Eating/Drinking',
        'Toileting',
        'Meal Preparation',
        'Medication Management',
        'Communication'
    ]

    for i, activity in enumerate(activities, 1):
        table.rows[i].cells[0].text = activity
        table.rows[i].cells[1].text = '[ ] Independent\n[ ] Supervision\n[ ] Assistance\n[ ] Dependent'
        table.rows[i].cells[2].text = ''

    doc.add_paragraph()

    # Psychosocial Needs
    doc.add_heading('PSYCHOSOCIAL & EMOTIONAL NEEDS', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    cell = table.rows[0].cells[0]
    cell.text = 'Cognitive status, mental health, social connections, hobbies, interests:\n\n\n\n\n'

    doc.add_paragraph()

    # Care Goals
    doc.add_heading('CARE GOALS', 1)
    for i in range(1, 4):
        doc.add_heading(f'Goal {i}:', 2)
        table = doc.add_table(rows=3, cols=2)
        add_table_border(table)

        goal_fields = [
            ('Goal:', '[Describe care goal]'),
            ('Actions/Interventions:', '[Steps to achieve goal]'),
            ('Target Date:', '[DD/MM/YYYY]')
        ]

        for j, (label, value) in enumerate(goal_fields):
            row = table.rows[j]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            set_cell_background(row.cells[0], 'E7E6E6')
            row.cells[1].text = value

        doc.add_paragraph()

    # Risk Assessment
    doc.add_heading('IDENTIFIED RISKS & MANAGEMENT STRATEGIES', 1)
    table = doc.add_table(rows=5, cols=3)
    add_table_border(table)

    headers = ['Risk Type', 'Risk Level', 'Management Strategy']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    risks = ['Falls', 'Medication', 'Nutrition', 'Infection']

    for i, risk in enumerate(risks, 1):
        table.rows[i].cells[0].text = risk
        table.rows[i].cells[1].text = '[ ] Low  [ ] Medium  [ ] High'
        table.rows[i].cells[2].text = ''

    doc.add_paragraph()

    # Services Provided
    doc.add_heading('CARE SERVICES TO BE PROVIDED', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    cell = table.rows[0].cells[0]
    cell.text = 'List all care services, frequency, and responsible staff:\n\n\n\n\n'

    doc.add_paragraph()

    # Signatures
    doc.add_heading('SIGNATURES', 1)
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Client/Representative Signature:'
    table.rows[0].cells[1].text = 'Care Coordinator Signature:'
    table.rows[1].cells[0].text = '\n\n'
    table.rows[1].cells[1].text = '\n\n'
    table.rows[2].cells[0].text = 'Date: [DD/MM/YYYY]'
    table.rows[2].cells[1].text = 'Date: [DD/MM/YYYY]'

    return doc

def create_ndis_timesheet():
    """Create NDIS Timesheet Template"""
    doc = Document()

    # Title
    title = doc.add_heading('NDIS SUPPORT WORKER TIMESHEET', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Period
    table = doc.add_table(rows=1, cols=2)
    add_table_border(table)
    table.rows[0].cells[0].text = 'Period:'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[0], 'E7E6E6')
    table.rows[0].cells[1].text = '[Start Date] to [End Date]'

    doc.add_paragraph()

    # Provider/Worker Details
    doc.add_heading('PROVIDER & WORKER DETAILS', 1)
    table = doc.add_table(rows=5, cols=2)
    add_table_border(table)

    fields = [
        ('Provider Name:', '[Provider Name]'),
        ('Provider ABN:', '[ABN]'),
        ('Worker Name:', '[Support Worker Name]'),
        ('Worker ID:', '[Employee/Contractor ID]'),
        ('Contact:', '[Phone/Email]')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Participant Details
    doc.add_heading('PARTICIPANT DETAILS', 1)
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    participant_fields = [
        ('Participant Name:', '[Participant Name]'),
        ('NDIS Number:', '[NDIS Number]'),
        ('Service Location:', '[Address where support provided]')
    ]

    for i, (label, value) in enumerate(participant_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Timesheet Entries
    doc.add_heading('HOURS WORKED', 1)
    table = doc.add_table(rows=15, cols=8)
    add_table_border(table)

    headers = ['Date', 'Day', 'Start Time', 'End Time', 'Break (mins)', 'Total Hours', 'Support Type', 'Participant Initials']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    days = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']

    for i in range(1, 15):
        for j in range(8):
            table.rows[i].cells[j].text = ''

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = 'TOTAL HOURS:'
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(total_row.cells[0], 'D9D9D9')
    total_row.cells[5].text = '[Total]'
    total_row.cells[5].paragraphs[0].runs[0].bold = True
    set_cell_background(total_row.cells[5], 'D9D9D9')

    doc.add_paragraph()

    # Support Item Numbers
    doc.add_heading('SUPPORT ITEM NUMBERS USED', 1)
    table = doc.add_table(rows=4, cols=3)
    add_table_border(table)

    headers = ['Support Type', 'Support Item Number', 'Total Hours']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 4):
        table.rows[i].cells[0].text = '[e.g., Daily Life Support - Weekday]'
        table.rows[i].cells[1].text = '[e.g., 01_011_0107_1_1]'
        table.rows[i].cells[2].text = ''

    doc.add_paragraph()

    # Travel/km
    doc.add_heading('TRAVEL/KILOMETRES (If applicable)', 1)
    table = doc.add_table(rows=2, cols=4)
    add_table_border(table)

    headers = ['Date', 'From/To', 'Kilometres', 'Purpose']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    doc.add_paragraph()

    # Signatures
    doc.add_heading('SIGNATURES', 1)
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Support Worker Signature:'
    table.rows[0].cells[1].text = 'Participant/Guardian Signature:'
    table.rows[1].cells[0].text = '\n\n'
    table.rows[1].cells[1].text = '\n\n'
    table.rows[2].cells[0].text = 'Date: [DD/MM/YYYY]'
    table.rows[2].cells[1].text = 'Date: [DD/MM/YYYY]'

    doc.add_paragraph()

    footer_note = doc.add_paragraph()
    footer_note.add_run('Note: ').bold = True
    footer_note.add_run('Both support worker and participant (or their representative) must sign this timesheet. Keep a copy for your records.')

    return doc

def create_progress_notes():
    """Create NDIS Progress Notes Template"""
    doc = Document()

    # Title
    title = doc.add_heading('NDIS SUPPORT WORKER PROGRESS NOTES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Session Details
    doc.add_heading('SESSION DETAILS', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)

    fields = [
        ('Date:', '[DD/MM/YYYY]'),
        ('Support Worker:', '[Name]'),
        ('Participant Name:', '[Participant Name]'),
        ('NDIS Number:', '[NDIS Number]'),
        ('Time:', '[Start] to [End]'),
        ('Location:', '[Where support provided]')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # SOAP Format
    doc.add_heading('SUBJECTIVE (What the participant said)', 1)
    doc.add_paragraph('Record the participant\'s description of their current state, feelings, concerns, or preferences:')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    doc.add_heading('OBJECTIVE (What you observed)', 1)
    doc.add_paragraph('Record your factual observations of the participant and the support session:')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    doc.add_heading('ASSESSMENT (Your professional assessment)', 1)
    doc.add_paragraph('Analysis of progress towards goals, changes in condition, effectiveness of supports:')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    doc.add_heading('PLAN (Next steps and actions)', 1)
    doc.add_paragraph('What will happen next, any changes needed, follow-up actions:')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    # Activities Completed
    doc.add_heading('ACTIVITIES COMPLETED', 1)
    activities = [
        '[ ] Personal care (showering, dressing, etc.)',
        '[ ] Meal preparation',
        '[ ] Household tasks',
        '[ ] Community access',
        '[ ] Transport',
        '[ ] Social engagement',
        '[ ] Skill development',
        '[ ] Exercise/physical activity',
        '[ ] Other: _______________'
    ]

    for activity in activities:
        doc.add_paragraph(activity, style='List Bullet')

    doc.add_paragraph()

    # Goal Progress
    doc.add_heading('PROGRESS TOWARDS NDIS GOALS', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    headers = ['Goal', 'Progress/Comments']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 4):
        table.rows[i].cells[0].text = f'Goal {i}:'
        table.rows[i].cells[1].text = ''

    doc.add_paragraph()

    # Incidents/Concerns
    doc.add_heading('INCIDENTS OR CONCERNS', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '[ ] No incidents or concerns to report\n[ ] Incident/concern: (describe below)\n\n\n'

    doc.add_paragraph()

    # Signature
    doc.add_heading('WORKER SIGNATURE', 1)
    table = doc.add_table(rows=2, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Signature:'
    table.rows[0].cells[1].text = 'Date:'
    table.rows[1].cells[0].text = '\n\n'
    table.rows[1].cells[1].text = '[DD/MM/YYYY]'

    return doc

def create_risk_assessment():
    """Create Aged Care Risk Assessment Template"""
    doc = Document()

    # Title
    title = doc.add_heading('AGED CARE RISK ASSESSMENT & MANAGEMENT PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Client Details
    doc.add_heading('CLIENT DETAILS', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    fields = [
        ('Client Name:', '[Full Name]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('Assessment Date:', '[DD/MM/YYYY]'),
        ('Assessor Name:', '[Staff Member Name]')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Risk Assessment Matrix
    doc.add_heading('RISK RATING GUIDE', 1)
    doc.add_paragraph('Low Risk: Unlikely to occur, minimal impact')
    doc.add_paragraph('Medium Risk: Possible to occur, moderate impact')
    doc.add_paragraph('High Risk: Likely to occur, significant impact')
    doc.add_paragraph('Critical Risk: Very likely or has occurred, severe impact')

    doc.add_paragraph()

    # 1. Falls Risk
    doc.add_heading('1. FALLS RISK ASSESSMENT', 1)
    table = doc.add_table(rows=11, cols=2)
    add_table_border(table)

    falls_factors = [
        ('History of falls in last 12 months', '[ ] Yes  [ ] No'),
        ('Use of mobility aid', '[ ] Yes  [ ] No  Type: _______'),
        ('Balance/gait issues', '[ ] Yes  [ ] No'),
        ('Dizziness or syncope', '[ ] Yes  [ ] No'),
        ('Medications affecting balance', '[ ] Yes  [ ] No'),
        ('Cognitive impairment', '[ ] Yes  [ ] No'),
        ('Vision/hearing impairment', '[ ] Yes  [ ] No'),
        ('Continence issues', '[ ] Yes  [ ] No'),
        ('Environmental hazards', '[ ] Yes  [ ] No'),
        ('Overall Falls Risk Rating', '[ ] Low  [ ] Medium  [ ] High  [ ] Critical')
    ]

    for i, (factor, response) in enumerate(falls_factors):
        row = table.rows[i]
        row.cells[0].text = factor
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = response

    doc.add_heading('Falls Management Strategies:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # 2. Medication Risk
    doc.add_heading('2. MEDICATION MANAGEMENT RISK', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    med_factors = [
        ('Number of medications', '_____ medications'),
        ('Self-administers medication', '[ ] Yes  [ ] No  [ ] With assistance'),
        ('Understands medication regime', '[ ] Yes  [ ] No  [ ] Partially'),
        ('Medication errors in last 6 months', '[ ] Yes  [ ] No'),
        ('High-risk medications (anticoagulants, insulin)', '[ ] Yes  [ ] No'),
        ('Overall Medication Risk Rating', '[ ] Low  [ ] Medium  [ ] High  [ ] Critical')
    ]

    for i, (factor, response) in enumerate(med_factors):
        row = table.rows[i]
        row.cells[0].text = factor
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = response

    doc.add_heading('Medication Management Strategies:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # 3. Nutrition Risk
    doc.add_heading('3. NUTRITION & HYDRATION RISK', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    nutrition_factors = [
        ('Recent unintentional weight loss', '[ ] Yes  [ ] No  Amount: _______'),
        ('Poor appetite or refusal to eat', '[ ] Yes  [ ] No'),
        ('Swallowing difficulties', '[ ] Yes  [ ] No'),
        ('Dehydration risk', '[ ] Yes  [ ] No'),
        ('Special diet requirements', '[ ] Yes  [ ] No  Type: _______'),
        ('Overall Nutrition Risk Rating', '[ ] Low  [ ] Medium  [ ] High  [ ] Critical')
    ]

    for i, (factor, response) in enumerate(nutrition_factors):
        row = table.rows[i]
        row.cells[0].text = factor
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = response

    doc.add_heading('Nutrition Management Strategies:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # 4. Infection Control Risk
    doc.add_heading('4. INFECTION CONTROL RISK', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    infection_factors = [
        ('Compromised immune system', '[ ] Yes  [ ] No'),
        ('Open wounds or skin integrity issues', '[ ] Yes  [ ] No'),
        ('Indwelling devices (catheter, PEG)', '[ ] Yes  [ ] No'),
        ('Recent or recurrent infections', '[ ] Yes  [ ] No'),
        ('Infection control compliance', '[ ] Good  [ ] Fair  [ ] Poor'),
        ('Overall Infection Risk Rating', '[ ] Low  [ ] Medium  [ ] High  [ ] Critical')
    ]

    for i, (factor, response) in enumerate(infection_factors):
        row = table.rows[i]
        row.cells[0].text = factor
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = response

    doc.add_heading('Infection Control Strategies:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # 5. Behavioral Risk
    doc.add_heading('5. BEHAVIORAL & PSYCHOLOGICAL RISK', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    behavioral_factors = [
        ('Wandering or absconding risk', '[ ] Yes  [ ] No'),
        ('Aggression or violence', '[ ] Yes  [ ] No'),
        ('Self-harm or suicidal ideation', '[ ] Yes  [ ] No'),
        ('Confusion or delirium', '[ ] Yes  [ ] No'),
        ('Non-compliance with care', '[ ] Yes  [ ] No'),
        ('Overall Behavioral Risk Rating', '[ ] Low  [ ] Medium  [ ] High  [ ] Critical')
    ]

    for i, (factor, response) in enumerate(behavioral_factors):
        row = table.rows[i]
        row.cells[0].text = factor
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = response

    doc.add_heading('Behavioral Management Strategies:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # Overall Summary
    doc.add_heading('OVERALL RISK SUMMARY', 1)
    table = doc.add_table(rows=2, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Overall Risk Level:'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[0], 'E7E6E6')
    table.rows[0].cells[1].text = '[ ] Low  [ ] Medium  [ ] High  [ ] Critical'

    table.rows[1].cells[0].text = 'Next Review Date:'
    table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[1].cells[0], 'E7E6E6')
    table.rows[1].cells[1].text = '[DD/MM/YYYY]'

    doc.add_paragraph()

    # Signatures
    doc.add_heading('SIGNATURES', 1)
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Assessor Signature:'
    table.rows[0].cells[1].text = 'Manager Review:'
    table.rows[1].cells[0].text = '\n\n'
    table.rows[1].cells[1].text = '\n\n'
    table.rows[2].cells[0].text = 'Date: [DD/MM/YYYY]'
    table.rows[2].cells[1].text = 'Date: [DD/MM/YYYY]'

    return doc

def create_medication_plan():
    """Create Medication Management Plan Template"""
    doc = Document()

    # Title
    title = doc.add_heading('MEDICATION MANAGEMENT PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Participant Details
    doc.add_heading('PARTICIPANT/CLIENT DETAILS', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    fields = [
        ('Name:', '[Full Name]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('NDIS Number / Client ID:', '[Number]'),
        ('Address:', '[Residential Address]'),
        ('GP Name:', '[Doctor Name]'),
        ('GP Contact:', '[Phone and Clinic]'),
        ('Pharmacy:', '[Pharmacy Name and Contact]')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Known Allergies
    doc.add_heading('KNOWN ALLERGIES & ADVERSE REACTIONS', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    set_cell_background(table.rows[0].cells[0], 'FFE6E6')
    table.rows[0].cells[0].text = '[ ] No Known Allergies\n[ ] Allergies: (list all medications, foods, and other allergens)\n\n\n'

    doc.add_paragraph()

    # Current Medications
    doc.add_heading('CURRENT MEDICATIONS', 1)
    table = doc.add_table(rows=11, cols=7)
    add_table_border(table)

    headers = ['Medication Name', 'Strength', 'Dose', 'Route', 'Time(s)', 'Prescriber', 'Purpose/Indication']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 11):
        for j in range(7):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # PRN Medications
    doc.add_heading('PRN (AS NEEDED) MEDICATIONS', 1)
    table = doc.add_table(rows=5, cols=6)
    add_table_border(table)

    headers = ['Medication Name', 'Strength/Dose', 'Route', 'Indication', 'Max Frequency', 'Prescriber']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 5):
        for j in range(6):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Administration Support
    doc.add_heading('MEDICATION ADMINISTRATION SUPPORT REQUIRED', 1)

    support_items = [
        '[ ] Self-administers independently',
        '[ ] Requires reminders/prompts',
        '[ ] Requires supervision',
        '[ ] Requires physical assistance',
        '[ ] Full administration by support worker',
        '[ ] Uses Webster pack/dose administration aid',
        '[ ] Requires assistance opening containers',
        '[ ] Other (specify): _______________'
    ]

    for item in support_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Special Instructions
    doc.add_heading('SPECIAL INSTRUCTIONS & CONSIDERATIONS', 1)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = 'Include timing with meals, crushing instructions, storage requirements, monitoring needs:\n\n\n\n\n'

    doc.add_paragraph()

    # Monitoring Requirements
    doc.add_heading('MONITORING REQUIREMENTS', 1)

    monitoring_items = [
        'Blood glucose monitoring: [ ] Yes  [ ] No   Frequency: _______',
        'Blood pressure monitoring: [ ] Yes  [ ] No   Frequency: _______',
        'Weight monitoring: [ ] Yes  [ ] No   Frequency: _______',
        'INR/anticoagulation monitoring: [ ] Yes  [ ] No   Frequency: _______',
        'Side effects to watch for: _______________',
        'Other monitoring: _______________'
    ]

    for item in monitoring_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Emergency Medications
    doc.add_heading('EMERGENCY MEDICATIONS', 1)
    table = doc.add_table(rows=3, cols=4)
    add_table_border(table)
    set_cell_background(table.rows[0].cells[0], 'FFE6E6')
    set_cell_background(table.rows[0].cells[1], 'FFE6E6')
    set_cell_background(table.rows[0].cells[2], 'FFE6E6')
    set_cell_background(table.rows[0].cells[3], 'FFE6E6')

    headers = ['Medication', 'Emergency Situation', 'Dose/Administration', 'Call Emergency Services?']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    examples = [
        ('e.g., EpiPen', 'Anaphylaxis/severe allergy', 'Auto-inject into thigh', '[ ] Yes'),
        ('e.g., Midazolam', 'Prolonged seizure', 'Buccal 0.5ml', '[ ] Yes if >5 mins')
    ]

    for i, (med, situation, admin, emergency) in enumerate(examples, 1):
        table.rows[i].cells[0].text = med
        table.rows[i].cells[1].text = situation
        table.rows[i].cells[2].text = admin
        table.rows[i].cells[3].text = emergency

    doc.add_paragraph()

    # Ceased Medications
    doc.add_heading('RECENTLY CEASED MEDICATIONS (Last 3 months)', 1)
    table = doc.add_table(rows=3, cols=4)
    add_table_border(table)

    headers = ['Medication Name', 'Date Ceased', 'Reason for Ceasing', 'Ceased By']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    for i in range(1, 3):
        for j in range(4):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Review Information
    doc.add_heading('MEDICATION PLAN REVIEW', 1)
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    review_fields = [
        ('Plan Created Date:', '[DD/MM/YYYY]'),
        ('Last Review Date:', '[DD/MM/YYYY]'),
        ('Next Review Due:', '[DD/MM/YYYY] (or when medications change)')
    ]

    for i, (label, value) in enumerate(review_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Signatures
    doc.add_heading('SIGNATURES', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Participant/Representative:'
    table.rows[0].cells[1].text = 'Care Coordinator/Manager:'
    table.rows[1].cells[0].text = 'Signature:\n\n'
    table.rows[1].cells[1].text = 'Signature:\n\n'
    table.rows[2].cells[0].text = 'Date:'
    table.rows[2].cells[1].text = 'Date:'
    table.rows[3].cells[0].text = 'Consented to support workers\nadministering medications:\n[ ] Yes  [ ] No'
    table.rows[3].cells[1].text = 'Verified by GP/Pharmacist:\n[ ] Yes  [ ] No'

    doc.add_paragraph()

    footer_note = doc.add_paragraph()
    footer_note.add_run('Note: ').bold = True
    footer_note.add_run('This plan must be reviewed at least annually or whenever medications change. Support workers must have appropriate medication administration training and delegation if administering medications.')

    return doc

def create_intake_form():
    """Create Participant Intake & Assessment Form"""
    doc = Document()

    # Title
    title = doc.add_heading('PARTICIPANT INTAKE & ASSESSMENT FORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Section 1: Personal Details
    doc.add_heading('1. PERSONAL DETAILS', 1)
    table = doc.add_table(rows=10, cols=2)
    add_table_border(table)

    personal_fields = [
        ('Full Legal Name:', '[Full Name]'),
        ('Preferred Name:', '[Preferred Name]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('Gender:', '[ ] Male  [ ] Female  [ ] Non-binary  [ ] Prefer not to say'),
        ('NDIS Number:', '[NDIS Number]'),
        ('Residential Address:', '[Full Address]'),
        ('Postal Address (if different):', '[Address]'),
        ('Phone (Mobile):', '[Mobile Number]'),
        ('Phone (Home):', '[Home Number]'),
        ('Email:', '[Email Address]')
    ]

    for i, (label, value) in enumerate(personal_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 2: Emergency Contacts
    doc.add_heading('2. EMERGENCY CONTACTS', 1)

    for i in range(1, 3):
        doc.add_heading(f'Emergency Contact {i}:', 2)
        table = doc.add_table(rows=4, cols=2)
        add_table_border(table)

        contact_fields = [
            ('Name:', '[Contact Name]'),
            ('Relationship:', '[Relationship]'),
            ('Phone:', '[Phone Number]'),
            ('Email:', '[Email]')
        ]

        for j, (label, value) in enumerate(contact_fields):
            row = table.rows[j]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            set_cell_background(row.cells[0], 'E7E6E6')
            row.cells[1].text = value

        doc.add_paragraph()

    # Section 3: NDIS Plan Details
    doc.add_heading('3. NDIS PLAN DETAILS', 1)
    table = doc.add_table(rows=5, cols=2)
    add_table_border(table)

    plan_fields = [
        ('Plan Start Date:', '[DD/MM/YYYY]'),
        ('Plan End Date:', '[DD/MM/YYYY]'),
        ('Plan Management Type:', '[ ] NDIA Managed  [ ] Plan Managed  [ ] Self Managed'),
        ('Plan Manager (if applicable):', '[Name and Contact]'),
        ('Support Coordinator:', '[Name and Contact]')
    ]

    for i, (label, value) in enumerate(plan_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 4: Disability & Health Information
    doc.add_heading('4. DISABILITY & HEALTH INFORMATION', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)

    health_fields = [
        ('Primary Disability:', '[Disability]'),
        ('Other Diagnoses:', '[List all relevant diagnoses]'),
        ('Current Medications:', '[List or attach medication chart]'),
        ('Known Allergies:', '[List all allergies]'),
        ('GP Details:', '[Name, Clinic, Phone]'),
        ('Medicare Number:', '[Medicare Number]')
    ]

    for i, (label, value) in enumerate(health_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 5: Communication Needs
    doc.add_heading('5. COMMUNICATION & CULTURAL NEEDS', 1)
    table = doc.add_table(rows=7, cols=2)
    add_table_border(table)

    comm_fields = [
        ('Preferred Language:', '[Language]'),
        ('Interpreter Required?', '[ ] Yes  [ ] No   Language: _______'),
        ('Preferred Communication Method:', '[ ] Verbal  [ ] Written  [ ] Sign Language  [ ] AAC Device\n[ ] Other: _______'),
        ('Vision Support Needs:', '[ ] None  [ ] Glasses  [ ] Blind/Low Vision\nDetails: _______'),
        ('Hearing Support Needs:', '[ ] None  [ ] Hearing Aid  [ ] Deaf/Hard of Hearing\nDetails: _______'),
        ('Cultural/Religious Considerations:', '[Any specific needs or preferences]'),
        ('LGBTQIA+ Considerations:', '[Any specific needs or preferences]')
    ]

    for i, (label, value) in enumerate(comm_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 6: Support Needs Assessment
    doc.add_heading('6. SUPPORT NEEDS ASSESSMENT', 1)
    doc.add_paragraph('Please indicate the level of support needed for each area:')
    doc.add_paragraph('I = Independent, S = Supervision, A = Assistance, D = Dependent')

    doc.add_paragraph()

    table = doc.add_table(rows=13, cols=3)
    add_table_border(table)

    headers = ['Activity', 'Support Level', 'Notes/Equipment Used']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    activities = [
        'Personal Care (showering, toileting)',
        'Dressing',
        'Mobility',
        'Transfers',
        'Eating/Drinking',
        'Meal Preparation',
        'Household Tasks',
        'Medication Management',
        'Money Management',
        'Transport',
        'Social Participation',
        'Communication'
    ]

    for i, activity in enumerate(activities, 1):
        table.rows[i].cells[0].text = activity
        table.rows[i].cells[1].text = '[ ] I  [ ] S  [ ] A  [ ] D'
        table.rows[i].cells[2].text = ''

    doc.add_paragraph()

    # Section 7: Services Requested
    doc.add_heading('7. SERVICES REQUESTED', 1)

    services = [
        '[ ] Assistance with Daily Living',
        '[ ] Community Access',
        '[ ] Social Support',
        '[ ] Support Coordination',
        '[ ] Therapy Services (specify type): _______',
        '[ ] Transport',
        '[ ] Household Tasks',
        '[ ] Meal Preparation',
        '[ ] Personal Care',
        '[ ] Other (specify): _______'
    ]

    for service in services:
        doc.add_paragraph(service, style='List Bullet')

    doc.add_paragraph()

    # Section 8: Goals
    doc.add_heading('8. PARTICIPANT GOALS', 1)
    doc.add_paragraph('What are your main goals for the next 12 months?')

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n\n\n'

    doc.add_paragraph()

    # Section 9: Preferences
    doc.add_heading('9. SUPPORT PREFERENCES', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)

    pref_fields = [
        ('Gender preference for support workers:', '[ ] Male  [ ] Female  [ ] No preference'),
        ('Age preference:', '[ ] Specific age range: _______  [ ] No preference'),
        ('Language preference:', '[Language]'),
        ('Days/times preferred for support:', '[Days and times]'),
        ('Activities/interests to include:', '[List interests]'),
        ('Anything to avoid:', '[List any dislikes or triggers]')
    ]

    for i, (label, value) in enumerate(pref_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 10: Risk Factors
    doc.add_heading('10. RISK FACTORS & SAFETY CONSIDERATIONS', 1)

    risks = [
        'Falls risk: [ ] Yes  [ ] No   Details: _______',
        'Wandering risk: [ ] Yes  [ ] No   Details: _______',
        'Behavioral concerns: [ ] Yes  [ ] No   Details: _______',
        'Self-harm risk: [ ] Yes  [ ] No   Details: _______',
        'Aggression risk: [ ] Yes  [ ] No   Details: _______',
        'Other safety concerns: _______'
    ]

    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')

    doc.add_paragraph()

    # Section 11: Consent
    doc.add_heading('11. CONSENT & PRIVACY', 1)

    consents = [
        '[ ] I consent to [Provider Name] collecting, using, and disclosing my personal information in accordance with privacy laws',
        '[ ] I consent to [Provider Name] sharing relevant information with my NDIS planner, support coordinator, and plan manager',
        '[ ] I consent to [Provider Name] reporting incidents to the NDIS Commission as required',
        '[ ] I consent to photography/video for progress documentation (not for marketing)',
        '[ ] I have received and understood the Privacy Policy',
        '[ ] I have received and understood the Service Agreement'
    ]

    for consent in consents:
        doc.add_paragraph(consent, style='List Bullet')

    doc.add_paragraph()

    # Section 12: Signatures
    doc.add_heading('12. SIGNATURES', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'PARTICIPANT (or Legal Representative/Guardian)'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[0], 'D9D9D9')
    table.rows[0].cells[1].text = 'PROVIDER REPRESENTATIVE'
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[1], 'D9D9D9')

    table.rows[1].cells[0].text = 'Name: [Print Name]'
    table.rows[1].cells[1].text = 'Name: [Print Name]'

    table.rows[2].cells[0].text = 'Signature:\n\n'
    table.rows[2].cells[1].text = 'Signature:\n\n'

    table.rows[3].cells[0].text = 'Date: [DD/MM/YYYY]'
    table.rows[3].cells[1].text = 'Date: [DD/MM/YYYY]'

    table.rows[4].cells[0].text = 'Relationship (if signing on behalf):\n[Relationship]'
    table.rows[4].cells[1].text = 'Position: [Position Title]'

    table.rows[5].cells[0].text = 'Authority to sign:\n[e.g., Guardian, Power of Attorney]'
    table.rows[5].cells[1].text = 'Contact: [Email/Phone]'

    doc.add_paragraph()

    footer_note = doc.add_paragraph()
    footer_note.add_run('Note: ').bold = True
    footer_note.add_run('This information will be kept confidential and used only for the purpose of providing NDIS supports. Please notify us immediately of any changes to your circumstances or contact details.')

    return doc

def create_daily_living_log():
    """Create NDIS Daily Living Log Template"""
    doc = Document()

    # Title
    title = doc.add_heading('NDIS DAILY LIVING LOG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Period
    table = doc.add_table(rows=1, cols=2)
    add_table_border(table)
    table.rows[0].cells[0].text = 'Week of:'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[0], 'E7E6E6')
    table.rows[0].cells[1].text = '[Start Date] to [End Date]'

    doc.add_paragraph()

    # Participant Details
    doc.add_heading('PARTICIPANT DETAILS', 1)
    table = doc.add_table(rows=2, cols=2)
    add_table_border(table)

    fields = [
        ('Participant Name:', '[Name]'),
        ('NDIS Number:', '[NDIS Number]')
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Legend
    doc.add_heading('ASSISTANCE LEVEL LEGEND', 1)
    doc.add_paragraph('I = Independent (no assistance needed)')
    doc.add_paragraph('P = Prompted (verbal reminder)')
    doc.add_paragraph('A = Assisted (physical help provided)')
    doc.add_paragraph('F = Full support (worker completed task)')
    doc.add_paragraph('N/A = Not applicable or not done')

    doc.add_paragraph()

    # Daily Activities Table
    doc.add_heading('DAILY PERSONAL CARE ACTIVITIES', 1)
    table = doc.add_table(rows=10, cols=8)
    add_table_border(table)

    # Headers
    headers = ['Activity', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    activities = [
        'Showering/Bathing',
        'Oral Care',
        'Hair Care',
        'Dressing',
        'Toileting',
        'Continence Management',
        'Skin Care',
        'Grooming/Shaving',
        'Nail Care'
    ]

    for i, activity in enumerate(activities, 1):
        table.rows[i].cells[0].text = activity
        for j in range(1, 8):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Meals Table
    doc.add_heading('MEALS & NUTRITION', 1)
    table = doc.add_table(rows=5, cols=8)
    add_table_border(table)

    headers = ['Meal', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    meals = ['Breakfast', 'Lunch', 'Dinner', 'Hydration (glasses of water)']

    for i, meal in enumerate(meals, 1):
        table.rows[i].cells[0].text = meal
        for j in range(1, 8):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Mobility & Transfers
    doc.add_heading('MOBILITY & TRANSFERS', 1)
    table = doc.add_table(rows=4, cols=8)
    add_table_border(table)

    headers = ['Activity', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    mobility = ['Walking/Moving Around', 'Bed Transfers', 'Chair Transfers']

    for i, activity in enumerate(mobility, 1):
        table.rows[i].cells[0].text = activity
        for j in range(1, 8):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Household Tasks
    doc.add_heading('HOUSEHOLD TASKS', 1)
    table = doc.add_table(rows=6, cols=8)
    add_table_border(table)

    headers = ['Task', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    tasks = ['Meal Preparation', 'Cleaning', 'Laundry', 'Shopping', 'Medication Management']

    for i, task in enumerate(tasks, 1):
        table.rows[i].cells[0].text = task
        for j in range(1, 8):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Social & Community Activities
    doc.add_heading('SOCIAL & COMMUNITY ACTIVITIES', 1)
    table = doc.add_table(rows=5, cols=8)
    add_table_border(table)

    headers = ['Activity', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    social = ['Social Interaction', 'Community Access', 'Recreation/Hobbies', 'Exercise/Physical Activity']

    for i, activity in enumerate(social, 1):
        table.rows[i].cells[0].text = activity
        for j in range(1, 8):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    # Daily Notes Section
    doc.add_heading('DAILY NOTES & OBSERVATIONS', 1)

    days_of_week = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']

    for day in days_of_week:
        doc.add_heading(f'{day}:', 2)
        table = doc.add_table(rows=1, cols=1)
        add_table_border(table)
        table.rows[0].cells[0].text = 'Notes (mood, participation, concerns, achievements):\n\n\n'
        doc.add_paragraph()

    # Weekly Summary
    doc.add_heading('WEEKLY SUMMARY', 1)

    doc.add_heading('Progress Towards Goals:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    doc.add_heading('Concerns or Issues:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    doc.add_heading('Recommendations:', 2)
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # Sign-off
    doc.add_heading('COMPLETED BY', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Name:'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[0], 'E7E6E6')
    table.rows[0].cells[1].text = '[Support Worker/Carer Name]'

    table.rows[1].cells[0].text = 'Role:'
    table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[1].cells[0], 'E7E6E6')
    table.rows[1].cells[1].text = '[Position]'

    table.rows[2].cells[0].text = 'Signature:'
    table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[2].cells[0], 'E7E6E6')
    table.rows[2].cells[1].text = '\n\n'

    table.rows[3].cells[0].text = 'Date:'
    table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[3].cells[0], 'E7E6E6')
    table.rows[3].cells[1].text = '[DD/MM/YYYY]'

    doc.add_paragraph()

    footer_note = doc.add_paragraph()
    footer_note.add_run('Note: ').bold = True
    footer_note.add_run('This log should be completed daily and reviewed weekly. Use it to track progress, identify patterns, and support plan reviews. Keep confidential and store securely.')

    return doc

def create_carer_impact_statement():
    """Create NDIS Carer Impact Statement Template"""
    doc = Document()

    # Title
    title = doc.add_heading('NDIS CARER IMPACT STATEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run('This statement describes how your caring role affects your daily life and wellbeing. It helps NDIS planners understand the support you provide and ensures the participant receives appropriate funding.').font.size = Pt(11)

    doc.add_paragraph()

    # Section 1: Your Details
    doc.add_heading('1. YOUR DETAILS (CARER)', 1)
    table = doc.add_table(rows=6, cols=2)
    add_table_border(table)

    carer_fields = [
        ('Your Name:', '[Full Name]'),
        ('Your Relationship to Participant:', '[e.g., Parent, Spouse, Sibling, Adult Child]'),
        ('Your Contact Number:', '[Phone Number]'),
        ('Your Email:', '[Email Address]'),
        ('Date of Statement:', '[DD/MM/YYYY]'),
        ('Do you live with the participant?', '[ ] Yes  [ ] No')
    ]

    for i, (label, value) in enumerate(carer_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 2: Participant Details
    doc.add_heading('2. PARTICIPANT DETAILS', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    participant_fields = [
        ('Participant Name:', '[Full Name]'),
        ('NDIS Number:', '[NDIS Number]'),
        ('Date of Birth:', '[DD/MM/YYYY]'),
        ('Primary Disability:', '[Disability/Diagnosis]')
    ]

    for i, (label, value) in enumerate(participant_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 3: Your Caring Role
    doc.add_heading('3. YOUR CARING ROLE', 1)
    doc.add_paragraph('Describe how long you have been a carer and how your caring role has evolved:')

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    # Section 4: Daily Caring Tasks
    doc.add_heading('4. DAILY CARING TASKS', 1)
    doc.add_paragraph('Describe the support you provide on a typical day. Include morning routines, mealtimes, personal care, supervision, and bedtime:')

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = 'Morning (e.g., waking, toileting, showering, dressing, breakfast, medications):\n\n\n'

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = 'Daytime (e.g., supervision, activities, meals, therapy appointments, behaviour support):\n\n\n'

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = 'Evening/Night (e.g., dinner, bathing, bedtime routine, overnight supervision):\n\n\n'

    doc.add_paragraph()

    # Section 5: Time Spent Caring
    doc.add_heading('5. TIME SPENT CARING', 1)
    table = doc.add_table(rows=4, cols=2)
    add_table_border(table)

    time_fields = [
        ('Hours of direct care per day:', '[e.g., 8-10 hours]'),
        ('Hours of supervision/monitoring per day:', '[e.g., 14 hours]'),
        ('Days per week you provide care:', '[e.g., 7 days]'),
        ('Do you need to wake at night?', '[ ] Yes  [ ] No   If yes, how often? ______')
    ]

    for i, (label, value) in enumerate(time_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 6: Impact on Your Life
    doc.add_heading('6. IMPACT ON YOUR LIFE', 1)

    doc.add_heading('Physical Health Impact:', 2)
    doc.add_paragraph('Describe any physical health impacts from your caring role (e.g., back pain, fatigue, injuries, missed medical appointments):')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    doc.add_heading('Mental and Emotional Health Impact:', 2)
    doc.add_paragraph('Describe any mental health or emotional impacts (e.g., stress, anxiety, depression, isolation, burnout):')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    doc.add_heading('Impact on Employment:', 2)
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    employment_fields = [
        ('Current employment status:', '[ ] Full-time  [ ] Part-time  [ ] Casual  [ ] Not working due to caring role  [ ] Retired'),
        ('Have you reduced work hours due to caring?', '[ ] Yes  [ ] No   If yes, by how many hours? ______'),
        ('Have you left work entirely due to caring?', '[ ] Yes  [ ] No   If yes, when? ______')
    ]

    for i, (label, value) in enumerate(employment_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(row.cells[0], 'E7E6E6')
        row.cells[1].text = value

    doc.add_paragraph()

    doc.add_heading('Impact on Relationships and Social Life:', 2)
    doc.add_paragraph('Describe how your caring role has affected your relationships with family, friends, and your ability to participate in social activities:')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    doc.add_heading('Impact on Other Family Members:', 2)
    doc.add_paragraph('Describe how the caring situation affects other family members (e.g., siblings, partner, other children):')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # Section 7: Sustainability
    doc.add_heading('7. SUSTAINABILITY OF YOUR CARING ROLE', 1)

    doc.add_paragraph('How sustainable is your current caring arrangement?')
    sustainability = [
        '[ ] I can continue caring with current supports',
        '[ ] I am struggling and need more support to continue',
        '[ ] My caring role is not sustainable without additional help',
        '[ ] I am at risk of burnout or breakdown'
    ]
    for item in sustainability:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('What would help you sustain your caring role?')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    # Section 8: What Happens Without Support
    doc.add_heading('8. WHAT HAPPENS IF YOU ARE UNABLE TO PROVIDE CARE?', 1)
    doc.add_paragraph('Describe what would happen to the participant if you were sick, injured, or unable to provide care:')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n'

    doc.add_paragraph()

    doc.add_paragraph('Who else can provide care if you are unavailable?')
    backup = [
        '[ ] No one - I am the sole carer',
        '[ ] Other family member (Name: ____________)',
        '[ ] Paid support worker',
        '[ ] Respite service'
    ]
    for item in backup:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Section 9: Goals and Aspirations
    doc.add_heading('9. YOUR GOALS AND ASPIRATIONS', 1)
    doc.add_paragraph('What are your goals for the future? What would you like to be able to do that your caring role currently prevents?')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    # Section 10: Additional Information
    doc.add_heading('10. ADDITIONAL INFORMATION', 1)
    doc.add_paragraph('Is there anything else you would like the NDIS planner to know about your caring role or its impact?')
    table = doc.add_table(rows=1, cols=1)
    add_table_border(table)
    table.rows[0].cells[0].text = '\n\n\n\n'

    doc.add_paragraph()

    # Declaration
    doc.add_heading('DECLARATION', 1)
    declaration = doc.add_paragraph()
    declaration.add_run('I declare that the information provided in this statement is true and accurate to the best of my knowledge. I understand this statement will be used to support the NDIS planning process for the participant named above.').font.size = Pt(11)

    doc.add_paragraph()

    # Signature
    table = doc.add_table(rows=3, cols=2)
    add_table_border(table)

    table.rows[0].cells[0].text = 'Carer Signature:'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[0].cells[0], 'E7E6E6')
    table.rows[0].cells[1].text = '\n\n'

    table.rows[1].cells[0].text = 'Print Name:'
    table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[1].cells[0], 'E7E6E6')
    table.rows[1].cells[1].text = '[Print Name]'

    table.rows[2].cells[0].text = 'Date:'
    table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(table.rows[2].cells[0], 'E7E6E6')
    table.rows[2].cells[1].text = '[DD/MM/YYYY]'

    doc.add_paragraph()

    # Footer note
    footer_note = doc.add_paragraph()
    footer_note.add_run('Tip: ').bold = True
    footer_note.add_run('Attach any supporting documentation such as medical letters, therapy reports, or daily care logs. Keep a copy of this statement for your records.')

    return doc

def create_schedule_of_supports_tracker():
    """Create NDIS Schedule of Supports Tracking Template (Excel)"""
    if not EXCEL_AVAILABLE:
        return None

    wb = Workbook()
    ws = wb.active
    ws.title = "Schedule of Supports"

    # Define styles
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    subheader_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    subheader_font = Font(bold=True, size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # Title and Instructions
    title_cell = ws['A1']
    title_cell.value = 'NDIS SCHEDULE OF SUPPORTS TRACKING TEMPLATE'
    title_cell.font = Font(bold=True, size=16, color="4472C4")
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:I1')

    instruction_cell = ws['A2']
    instruction_cell.value = 'Use this template to track your NDIS funding throughout your plan period. Update after each service or invoice to monitor spending.'
    instruction_cell.font = Font(size=10, italic=True)
    instruction_cell.alignment = Alignment(horizontal='center', wrap_text=True)
    ws.row_dimensions[2].height = 30
    ws.merge_cells('A2:I2')

    # Participant Details Section
    ws['A4'] = 'Participant Name:'
    ws['A4'].font = Font(bold=True)
    ws.merge_cells('A4:B4')

    ws['C4'] = '[Enter Name]'
    ws.merge_cells('C4:D4')

    ws['E4'] = 'NDIS Number:'
    ws['E4'].font = Font(bold=True)
    ws.merge_cells('E4:F4')

    ws['G4'] = '[Enter NDIS Number]'
    ws.merge_cells('G4:I4')

    ws['A5'] = 'Plan Start Date:'
    ws['A5'].font = Font(bold=True)
    ws.merge_cells('A5:B5')

    ws['C5'] = '[DD/MM/YYYY]'
    ws.merge_cells('C5:D5')

    ws['E5'] = 'Plan End Date:'
    ws['E5'].font = Font(bold=True)
    ws.merge_cells('E5:F5')

    ws['G5'] = '[DD/MM/YYYY]'
    ws.merge_cells('G5:I5')

    # Column Headers (Row 7)
    headers = [
        'Category',
        'Support Item',
        'Funded Units',
        'Price/Unit',
        'Total Budget',
        'Used',
        'Remaining',
        '% Used',
        'Notes'
    ]

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=7, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = border

    # Set column widths
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 12
    ws.column_dimensions['G'].width = 12
    ws.column_dimensions['H'].width = 10
    ws.column_dimensions['I'].width = 30

    # Sample data rows with categories
    sample_data = [
        # Core Supports
        ('CORE SUPPORTS', '', '', '', '', '', '', '', ''),
        ('Core - Daily Life', 'Personal Care', '200 hrs', '$67.56', '=C9*D9', '50 hrs', '=C9-F9', '=F9/C9', 'Review Q3'),
        ('Core - Daily Life', 'Domestic Assistance', '100 hrs', '$67.56', '=C10*D10', '25 hrs', '=C10-F10', '=F10/C10', 'On track'),
        ('Core - Community', 'Community Access', '100 hrs', '$67.56', '=C11*D11', '20 hrs', '=C11-F11', '=F11/C11', ''),
        ('Core - Transport', 'Transport', 'Flexible', '$1/km', '$2,000', '$500', '=E12-F12', '=F12/E12', ''),
        ('', '', '', '', '', '', '', '', ''),
        # Capacity Building
        ('CAPACITY BUILDING', '', '', '', '', '', '', '', ''),
        ('CB - Therapy', 'Occupational Therapy', '20 hrs', '$193.99', '=C15*D15', '8 hrs', '=C15-F15', '=F15/C15', ''),
        ('CB - Therapy', 'Physiotherapy', '15 hrs', '$193.99', '=C16*D16', '5 hrs', '=C16-F16', '=F16/C16', ''),
        ('CB - Support Coord', 'Support Coordination L1', '40 hrs', '$80.06', '=C17*D17', '15 hrs', '=C17-F17', '=F17/C17', ''),
        ('', '', '', '', '', '', '', '', ''),
        # Capital
        ('CAPITAL SUPPORTS', '', '', '', '', '', '', '', ''),
        ('Capital - AT', 'Wheelchair', '1', '$8,000', '$8,000', '$8,000', '=E20-F20', '=F20/E20', 'Purchased'),
        ('Capital - Home Mods', 'Bathroom Modifications', '1', '$12,000', '$12,000', '$0', '=E21-F21', '=F21/E21', 'Quote pending'),
    ]

    current_row = 8
    for row_data in sample_data:
        # Check if this is a category header row
        is_category_header = row_data[0] and 'SUPPORTS' in str(row_data[0]) and not row_data[1]

        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=current_row, column=col_num)
            cell.value = value
            cell.border = border

            # Style category headers
            if is_category_header:
                cell.fill = subheader_fill
                cell.font = subheader_font

            # Center align numeric columns
            if col_num in [3, 4, 6, 7, 8]:
                cell.alignment = Alignment(horizontal='center', vertical='center')
            elif col_num == 9:
                cell.alignment = Alignment(wrap_text=True, vertical='center')

        # Merge category header rows after setting all values
        if is_category_header:
            ws.merge_cells(f'A{current_row}:I{current_row}')

        current_row += 1

    # Totals Row
    totals_row = current_row + 1
    totals_cell = ws[f'A{totals_row}']
    totals_cell.value = 'TOTAL FUNDING:'
    totals_cell.font = Font(bold=True, size=12)
    totals_cell.alignment = Alignment(horizontal='right', vertical='center')
    totals_cell.fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
    ws.merge_cells(f'A{totals_row}:D{totals_row}')

    # Total budget formula (sum all budget cells)
    total_budget_cell = ws[f'E{totals_row}']
    total_budget_cell.value = f'=SUM(E9:E21)'
    total_budget_cell.font = Font(bold=True, size=12)
    total_budget_cell.fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
    total_budget_cell.border = border

    for col in ['F', 'G', 'H', 'I']:
        cell = ws[f'{col}{totals_row}']
        cell.fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
        cell.border = border

    # Add helpful tips section
    tips_row = totals_row + 3
    tips_title = ws[f'A{tips_row}']
    tips_title.value = 'TRACKING TIPS:'
    tips_title.font = Font(bold=True, size=11, color="4472C4")
    ws.merge_cells(f'A{tips_row}:I{tips_row}')

    tips = [
        "1. Update the 'Used' column after each service or when you receive invoices",
        "2. The 'Remaining' and '% Used' columns calculate automatically",
        "3. Aim to use approximately 50% of each budget at the halfway point of your plan",
        "4. Use the 'Notes' column to track important information (e.g., 'booking next month', 'waiting for quote')",
        "5. Review this tracker monthly with your Support Coordinator",
        "6. Unused funding does not roll over - use it or lose it!",
        "7. If you're running low on funding early, request a plan review"
    ]

    for i, tip in enumerate(tips):
        tip_row = tips_row + i + 1
        tip_cell = ws[f'A{tip_row}']
        tip_cell.value = tip
        tip_cell.font = Font(size=10)
        tip_cell.alignment = Alignment(wrap_text=True)
        ws.row_dimensions[tip_row].height = 18
        ws.merge_cells(f'A{tip_row}:I{tip_row}')

    # Format percentage column
    for row in range(9, current_row):
        cell = ws[f'H{row}']
        if cell.value and str(cell.value).startswith('='):
            cell.number_format = '0%'

    # Format currency columns
    for row in range(9, current_row):
        for col in ['D', 'E']:
            cell = ws[f'{col}{row}']
            if cell.value and (str(cell.value).startswith('$') or str(cell.value).startswith('=')):
                cell.number_format = '$#,##0.00'

    # Freeze panes at header row
    ws.freeze_panes = 'A8'

    return wb

def main():
    """Main function to create all templates"""
    output_dir = '/Users/andredeansmith/mdhomecarebuild/public/downloads'

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # DOCX templates
    docx_templates = {
        'ndis-service-agreement-template-2026.docx': create_ndis_service_agreement,
        'ndis-incident-report-form-2026.docx': create_ndis_incident_report,
        'ndis-support-plan-template-2026.docx': create_ndis_support_plan,
        'aged-care-individual-care-plan-2026.docx': create_aged_care_plan,
        'ndis-timesheet-template-2026.docx': create_ndis_timesheet,
        'ndis-progress-notes-template-2026.docx': create_progress_notes,
        'aged-care-risk-assessment-template-2026.docx': create_risk_assessment,
        'medication-management-plan-template-2026.docx': create_medication_plan,
        'ndis-participant-intake-form-2026.docx': create_intake_form,
        'ndis-daily-living-log-template-2026.docx': create_daily_living_log,
        'ndis-carer-impact-statement-template-2026.docx': create_carer_impact_statement
    }

    # Excel templates
    xlsx_templates = {}
    if EXCEL_AVAILABLE:
        xlsx_templates = {
            'ndis-schedule-supports-template-2026.xlsx': create_schedule_of_supports_tracker
        }

    total_templates = len(docx_templates) + len(xlsx_templates)
    print(f'Creating {total_templates} NDIS/Aged Care templates...\n')

    # Create DOCX templates
    for filename, create_func in docx_templates.items():
        try:
            print(f'Creating {filename}...')
            doc = create_func()
            filepath = os.path.join(output_dir, filename)
            doc.save(filepath)
            print(f'✓ Created: {filepath}\n')
        except Exception as e:
            print(f'✗ Error creating {filename}: {e}\n')

    # Create XLSX templates
    for filename, create_func in xlsx_templates.items():
        try:
            print(f'Creating {filename}...')
            wb = create_func()
            if wb:
                filepath = os.path.join(output_dir, filename)
                wb.save(filepath)
                print(f'✓ Created: {filepath}\n')
            else:
                print(f'⊘ Skipped {filename} (openpyxl not available)\n')
        except Exception as e:
            print(f'✗ Error creating {filename}: {e}\n')

    print(f'Template creation complete! All files saved to {output_dir}')

if __name__ == '__main__':
    main()
